from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
IMG = ROOT / "tmp" / "quiz_images"
OUT.mkdir(exist_ok=True)
IMG.mkdir(parents=True, exist_ok=True)

GREEN = "176B52"
DARK = "17332D"
GOLD = "E4A93A"
MINT = "EAF5F0"
PALE_GOLD = "FFF6DE"
GRAY = "5A6763"
LIGHT = "F5F7F6"
WHITE = "FFFFFF"
RED = "A3423C"


topics = [
    {
        "title": "1. Bioenergetics: The Cell's Energy Budget",
        "tag": "ENERGY LAB",
        "scenario": "A campus food-science team is designing a rapid test for whether cellular reactions can proceed. They must distinguish energy-releasing reactions from energy-requiring ones and decide when ATP coupling can make an otherwise unfavourable process proceed.",
        "image": "bioenergetics.png",
        "questions": [
            {"n": 1, "level": "Remember", "marks": 1,
             "a": "State what bioenergetics studies and identify the sign of Delta G for a spontaneous reaction.",
             "b": "State what Gibbs free-energy change predicts and identify the sign of Delta G at equilibrium.",
             "ans_a": "Bioenergetics studies energy flow and transformation in living systems; a spontaneous reaction has Delta G < 0.",
             "ans_b": "Delta G predicts reaction spontaneity/direction; at equilibrium Delta G = 0."},
            {"n": 2, "level": "Understand", "marks": 2,
             "a": "Why can an anabolic reaction with a positive Delta G proceed when it is coupled to ATP hydrolysis? Answer in two sentences.",
             "b": "Explain why calling ATP an 'energy currency' does not mean ATP creates energy.",
             "ans_a": "ATP hydrolysis has a sufficiently negative Delta G. When the reactions are mechanistically coupled, their Delta G values add; the overall process can become negative.",
             "ans_b": "ATP transfers conserved chemical free energy from energy-yielding processes to energy-requiring work. It transforms and couples energy rather than creating it."},
            {"n": 3, "level": "Apply", "marks": 3,
             "a": "A biosynthetic step has Delta G = +18.0 kJ/mol and is coupled to ATP hydrolysis with Delta G = -30.5 kJ/mol. Calculate the combined Delta G and predict whether the coupled process is thermodynamically favourable.",
             "b": "Transport of a solute has Delta G = +22.0 kJ/mol. It is coupled to a reaction with Delta G = -28.0 kJ/mol. Calculate the net Delta G and state what would happen if coupling were lost.",
             "ans_a": "Delta G(total) = +18.0 - 30.5 = -12.5 kJ/mol; the coupled process is favourable/spontaneous.",
             "ans_b": "Delta G(total) = +22.0 - 28.0 = -6.0 kJ/mol; coupled transport is favourable. Without coupling, the +22.0 kJ/mol transport step would be unfavourable."},
        ],
    },
    {
        "title": "2. Glycolysis: The Ten-Step Sprint",
        "tag": "SPRINT TRACK",
        "scenario": "During a short, intense sprint, skeletal muscle needs ATP faster than oxygen delivery can support oxidative metabolism. Glycolysis must keep running, and its irreversible steps act as metabolic control gates.",
        "image": "glycolysis.png",
        "questions": [
            {"n": 4, "level": "Remember", "marks": 1,
             "a": "Give the net products of glycolysis from one glucose molecule under standard accounting.",
             "b": "Name the ATP-investment and ATP-payoff amounts in glycolysis and calculate the net ATP yield.",
             "ans_a": "2 pyruvate, 2 ATP net, 2 NADH, and 2 H2O (plus 2 H+ depending on convention).",
             "ans_b": "2 ATP are invested and 4 ATP are produced; net yield = 2 ATP per glucose."},
            {"n": 5, "level": "Understand", "marks": 2,
             "a": "Explain why the PFK-1 reaction is called the first committed step of glycolysis. Name one activator and one inhibitor from the lectures.",
             "b": "Differentiate the roles of hexokinase/glucokinase and PFK-1 in controlling glucose use. Include one regulatory feature of each.",
             "ans_a": "PFK-1 commits fructose-6-phosphate to glycolysis by forming fructose-1,6-bisphosphate. Activators include AMP or fructose-2,6-bisphosphate; inhibitors include ATP, citrate, low pH, or fatty acids.",
             "ans_b": "Hexokinase/glucokinase traps glucose as G6P; hexokinase is product-inhibited by G6P, while liver glucokinase has lower affinity and is not inhibited by G6P. PFK-1 controls the committed/rate-limiting step and responds to the energy state."},
            {"n": 6, "level": "Apply", "marks": 3,
             "a": "Use Figure 2. During oxygen-limited sprinting, explain why pyruvate is converted to lactate and predict the immediate effect on glycolytic ATP production if lactate dehydrogenase is blocked.",
             "b": "A red blood cell cannot use mitochondria. Explain why lactate production is essential for its continued ATP supply and identify the cofactor that is regenerated.",
             "ans_a": "Lactate dehydrogenase oxidizes NADH to NAD+, allowing GAPDH and glycolysis to continue. Blocking it depletes cytosolic NAD+, slows/stops glycolysis, and rapidly lowers ATP production.",
             "ans_b": "RBCs depend on glycolysis for ATP. Conversion of pyruvate to lactate regenerates NAD+, which is required by glyceraldehyde-3-phosphate dehydrogenase."},
        ],
    },
    {
        "title": "3. Glycogen Metabolism: The Branched Emergency Store",
        "tag": "FUEL DEPOT",
        "scenario": "A student athlete alternates between lectures, training, and an overnight fast. Liver glycogen supports blood glucose, while muscle glycogen is a rapid local fuel. The highly branched structure determines how quickly glucose units can be released.",
        "image": "glycogen.png",
        "questions": [
            {"n": 7, "level": "Remember", "marks": 1,
             "a": "Identify the glycosidic bond in the linear chains and the bond at glycogen branch points.",
             "b": "What is a non-reducing end, and why do many non-reducing ends help rapid glycogen mobilization?",
             "ans_a": "Linear chains contain alpha(1->4) bonds; branch points contain alpha(1->6) bonds.",
             "ans_b": "A non-reducing end is a chain end from which glycogen-metabolizing enzymes add or remove glucose. Many ends allow many enzyme molecules to work simultaneously, increasing the rate."},
            {"n": 8, "level": "Understand", "marks": 2,
             "a": "Why can liver glycogen help maintain blood glucose, whereas muscle glycogen mainly supplies the muscle itself?",
             "b": "After glycogen phosphorylase releases glucose-1-phosphate, trace its conversion in liver until free glucose can enter the blood.",
             "ans_a": "Liver expresses glucose-6-phosphatase and can convert G6P to free glucose for export. Muscle lacks glucose-6-phosphatase, so G6P enters glycolysis locally.",
             "ans_b": "G1P is converted by phosphoglucomutase to G6P; it enters the ER and glucose-6-phosphatase hydrolyses it to glucose, which is returned to the cytosol and released."},
            {"n": 9, "level": "Apply", "marks": 3,
             "a": "A patient has normal glycogen phosphorylase but deficient debranching enzyme. Predict the glycogen structure that accumulates and explain why breakdown stops near branches.",
             "b": "A muscle sample lacks glycogen phosphorylase activity but has a normal debranching enzyme. Predict the effect on exercise tolerance and on release of glucose-1-phosphate from alpha(1->4) bonds.",
             "ans_a": "Short outer chains/limit-dextrin-like glycogen accumulates. Phosphorylase stops about four residues from a branch and cannot cleave alpha(1->6); transferase and alpha-1,6-glucosidase activities are needed.",
             "ans_b": "Rapid glycogen use and exercise tolerance decrease. Debranching enzyme cannot substitute for phosphorylase, so phosphorolytic release of G1P from alpha(1->4) bonds is severely impaired."},
        ],
    },
    {
        "title": "4. Gluconeogenesis: Rebuilding Glucose During a Fast",
        "tag": "FASTING RESCUE",
        "scenario": "After an overnight fast, hepatic glycogen becomes limited. The liver uses lactate, glycerol, and glucogenic amino-acid carbon skeletons to rebuild glucose, paying an energy cost to bypass the irreversible steps of glycolysis.",
        "image": "gluconeogenesis.png",
        "questions": [
            {"n": 10, "level": "Remember", "marks": 1,
             "a": "Name three gluconeogenic precursors and the main organ that performs gluconeogenesis.",
             "b": "Name two tissues with a continuous need for glucose and state why gluconeogenesis becomes important after prolonged fasting.",
             "ans_a": "Any three: lactate, glycerol, pyruvate, alanine/other glucogenic amino acids, TCA intermediates, propionate. Main organ: liver (kidney contributes during prolonged fasting).",
             "ans_b": "Examples include brain, RBCs, lens, adrenal medulla, and exercising muscle. As hepatic glycogen is depleted, gluconeogenesis becomes the major source maintaining blood glucose."},
            {"n": 11, "level": "Understand", "marks": 2,
             "a": "Match each irreversible glycolytic enzyme with its gluconeogenic bypass enzyme(s): hexokinase, PFK-1, and pyruvate kinase.",
             "b": "Why is oxaloacetate commonly converted to malate before its carbon skeleton moves from mitochondrion to cytosol during gluconeogenesis?",
             "ans_a": "Hexokinase/glucokinase -> glucose-6-phosphatase; PFK-1 -> fructose-1,6-bisphosphatase; pyruvate kinase -> pyruvate carboxylase plus PEP carboxykinase.",
             "ans_b": "The inner mitochondrial membrane is poorly permeable to oxaloacetate. Malate can be transported and then reoxidized to oxaloacetate in the cytosol, also transferring reducing equivalents."},
            {"n": 12, "level": "Apply", "marks": 3,
             "a": "Calculate the high-energy phosphate cost of producing one glucose from two pyruvate. Separate ATP and GTP, and include the NADH requirement.",
             "b": "Two glucose molecules must be synthesized from pyruvate. Calculate the total ATP, GTP, and NADH required, using the lecture stoichiometry.",
             "ans_a": "4 ATP + 2 GTP = 6 high-energy phosphate equivalents, plus 2 NADH per glucose.",
             "ans_b": "8 ATP, 4 GTP, and 4 NADH for two glucose molecules; 12 high-energy phosphate equivalents in total."},
        ],
    },
    {
        "title": "5. Reciprocal Regulation: Choosing Store, Burn, or Build",
        "tag": "CONTROL ROOM",
        "scenario": "A hepatocyte receives different hormonal instructions after lunch, during an overnight fast, and during acute stress. Reciprocal regulation prevents glycolysis and gluconeogenesis - or glycogenesis and glycogenolysis - from running strongly at the same time.",
        "image": "regulation.png",
        "questions": [
            {"n": 13, "level": "Remember", "marks": 1,
             "a": "State the main effect of insulin and of glucagon on liver glycogen metabolism.",
             "b": "State the main glycogen effect of epinephrine and name the physiological situation that triggers it.",
             "ans_a": "Insulin promotes glycogen synthesis; glucagon promotes liver glycogen breakdown and inhibits synthesis.",
             "ans_b": "Epinephrine promotes rapid glycogen breakdown during acute stress or exercise."},
            {"n": 14, "level": "Understand", "marks": 2,
             "a": "Explain the reciprocal effect of fructose-2,6-bisphosphate on PFK-1 and fructose-1,6-bisphosphatase.",
             "b": "Why does a fall in hepatic fructose-2,6-bisphosphate during fasting reduce futile cycling between glycolysis and gluconeogenesis?",
             "ans_a": "Fructose-2,6-bisphosphate activates PFK-1, stimulating glycolysis, and inhibits FBPase-1, suppressing gluconeogenesis.",
             "ans_b": "Lower F2,6BP removes activation of PFK-1 and relieves inhibition of FBPase-1, coordinating lower glycolysis with higher gluconeogenesis instead of running both strongly."},
            {"n": 15, "level": "Analyze", "marks": 3,
             "a": "Use Figure 5. During fasting, trace the sequence from glucagon to PKA to the bifunctional PFK-2/FBPase-2 enzyme. Predict F2,6BP, glycolysis, and gluconeogenesis.",
             "b": "After a carbohydrate-rich meal, trace insulin action through dephosphorylation of hepatic PFK-2/FBPase-2. Predict F2,6BP and the direction of glucose metabolism.",
             "ans_a": "Glucagon -> cAMP -> PKA -> phosphorylation: PFK-2 decreases and FBPase-2 increases -> F2,6BP falls -> glycolysis decreases and gluconeogenesis increases.",
             "ans_b": "Insulin promotes phosphatase action/dephosphorylation: PFK-2 increases and FBPase-2 decreases -> F2,6BP rises -> glycolysis increases and gluconeogenesis decreases; storage is favoured."},
        ],
    },
    {
        "title": "6. Pentose Phosphate Pathway: Reducing Power and Ribose",
        "tag": "ANTIOXIDANT SHIELD",
        "scenario": "Red blood cells face oxidants but have no mitochondria. Other cells may need NADPH for reductive biosynthesis or ribose-5-phosphate for nucleotide synthesis. The pentose phosphate pathway adjusts its oxidative and non-oxidative phases to those needs.",
        "image": "ppp.png",
        "questions": [
            {"n": 16, "level": "Understand", "marks": 2,
             "a": "Compare the main products and reversibility of the oxidative and non-oxidative phases of the pentose phosphate pathway.",
             "b": "A rapidly dividing cell needs ribose-5-phosphate more than NADPH. Explain how reversible non-oxidative reactions can help meet this need using glycolytic intermediates.",
             "ans_a": "The oxidative phase is irreversible and produces NADPH, CO2, and ribulose-5-phosphate. The non-oxidative phase is reversible and interconverts pentose phosphates with fructose-6-phosphate and glyceraldehyde-3-phosphate.",
             "ans_b": "Reversible transketolase/transaldolase reactions can run from F6P and G3P toward ribose-5-phosphate, allowing ribose production without proportionally high oxidative-phase NADPH production."},
            {"n": 17, "level": "Analyze", "marks": 4,
             "a": "A patient with low glucose-6-phosphate dehydrogenase activity develops haemolysis after oxidant-drug exposure. Build the causal chain from reduced PPP flux to red-cell membrane damage.",
             "b": "Two samples are exposed to peroxide: normal RBCs and RBCs with impaired oxidative PPP activity. Predict which has the lower reduced-glutathione level and explain the roles of NADPH and glutathione reductase.",
             "ans_a": "Low G6PD -> less oxidative PPP NADPH -> reduced ability of glutathione reductase to regenerate GSH -> peroxide/ROS accumulate -> haemoglobin and membrane proteins/lipids are oxidized -> RBC damage and haemolysis.",
             "ans_b": "The impaired-PPP RBCs have lower GSH. NADPH donates reducing power to glutathione reductase, which converts oxidized glutathione (GSSG) back to GSH for peroxide detoxification."},
        ],
    },
    {
        "title": "7. CAC and Glyoxylate Cycle: The Carbon-Routing Finale",
        "tag": "CARBON ROUNDABOUT",
        "scenario": "A germinating oil seed converts stored fat into carbohydrate, while a mammalian liver cannot obtain net glucose from even-chain fatty acids. Both use acetyl-CoA, but only the seed can bypass the carbon-losing steps of the citric acid cycle through the glyoxylate cycle.",
        "image": "cac_glyoxylate.png",
        "questions": [
            {"n": 18, "level": "Analyze", "marks": 4,
             "a": "For one acetyl-CoA entering the CAC, account for NADH, FADH2, GTP/ATP, and CO2. Then explain how high ATP and NADH affect cycle flux.",
             "b": "A cell has high NADH and succinyl-CoA. Identify two CAC control points likely to slow and explain the product/energy feedback involved.",
             "ans_a": "Per acetyl-CoA: 3 NADH, 1 FADH2, 1 GTP (or ATP), and 2 CO2; about 10 ATP equivalents after oxidative phosphorylation. High ATP/NADH signal energy sufficiency and inhibit major control enzymes, reducing flux.",
             "ans_b": "Isocitrate dehydrogenase is inhibited by NADH; alpha-ketoglutarate dehydrogenase is inhibited by NADH and succinyl-CoA. Citrate synthase/PDH may also slow under high-energy/product conditions."},
            {"n": 19, "level": "Evaluate", "marks": 5,
             "a": "Defend this statement: 'A germinating seed can achieve net carbohydrate synthesis from acetyl-CoA, but a vertebrate cannot.' Refer to carbon loss and the two unique glyoxylate-cycle enzymes.",
             "b": "A researcher proposes inserting only isocitrate lyase into a mammalian cell to create a functional glyoxylate cycle. Evaluate the proposal and identify the additional enzyme and carbon-routing requirement.",
             "ans_a": "The CAC loses two carbons as CO2, preventing net oxaloacetate gain from acetyl-CoA. The glyoxylate cycle bypasses the decarboxylations: isocitrate lyase forms succinate + glyoxylate, and malate synthase condenses glyoxylate with a second acetyl-CoA. Succinate can support carbohydrate synthesis. Vertebrates lack this pathway.",
             "ans_b": "Isocitrate lyase alone is insufficient. Malate synthase is also needed to combine glyoxylate with acetyl-CoA, and the pathway must route isocitrate away from CAC decarboxylations while regenerating oxaloacetate and exporting a four-carbon product such as succinate."},
            {"n": 20, "level": "Create", "marks": 5,
             "a": "Design a carbon-flow plan for a fed liver cell with high ATP, abundant glucose, demand for fatty-acid synthesis, and demand for nucleotides. Route glucose among glycolysis/CAC, glycogen, and PPP; justify each choice using energy and biosynthetic signals.",
             "b": "Design a carbon-flow plan for a fasting liver cell with low blood glucose and active fatty-acid oxidation. Integrate glycogenolysis, gluconeogenesis, glycolysis regulation, the CAC, and the fact that acetyl-CoA cannot give net glucose in vertebrates.",
             "ans_a": "A strong answer routes glucose to glycogen storage; uses glycolysis to supply pyruvate/acetyl-CoA and citrate for fatty-acid synthesis but slows complete oxidation when ATP/NADH are high; directs G6P through oxidative PPP for NADPH and through non-oxidative PPP for ribose-5-P. It should mention insulin/F2,6BP favouring glycolysis and suppressing gluconeogenesis.",
             "ans_b": "A strong answer uses early liver glycogenolysis and then gluconeogenesis from lactate, glycerol, and glucogenic amino acids; glucagon lowers F2,6BP, suppressing glycolysis and activating gluconeogenesis. Fatty-acid oxidation supplies ATP/NADH and acetyl-CoA; acetyl-CoA supports energy/ketone production and activates pyruvate carboxylase but cannot provide net glucose because vertebrates lack the glyoxylate cycle."},
        ],
    },
]


CANVAS = (1710, 576)
FONT_REG = r"C:\Windows\Fonts\arial.ttf"
FONT_BOLD = r"C:\Windows\Fonts\arialbd.ttf"


def font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)


def xy(frac):
    return (int(frac[0]*CANVAS[0]), int(frac[1]*CANVAS[1]))


def new_canvas(title):
    im=Image.new("RGB",CANVAS,"white"); d=ImageDraw.Draw(im)
    d.text((34,22),title,font=font(34,True),fill="#"+DARK)
    return im,d


def pbox(d,x,y,w,h,text,fc=MINT,ec=GREEN,size=23):
    x1,y1=xy((x,y)); x2,y2=xy((x+w,y+h))
    d.rounded_rectangle((x1,y1,x2,y2),radius=18,fill="#"+fc,outline="#"+ec,width=4)
    lines=text.split("\n"); f=font(size,True); heights=[]
    for line in lines: heights.append(d.textbbox((0,0),line,font=f)[3])
    total=sum(heights)+8*(len(lines)-1); yy=(y1+y2-total)//2
    for line,hh in zip(lines,heights):
        bb=d.textbbox((0,0),line,font=f); xx=(x1+x2-(bb[2]-bb[0]))//2
        d.text((xx,yy),line,font=f,fill="#"+DARK); yy+=hh+8


def arrow(d,a,b,color=GREEN,both=False,width=6):
    a=xy(a); b=xy(b); d.line((a,b),fill="#"+color,width=width)
    def head(tip,tail):
        ang=math.atan2(tip[1]-tail[1],tip[0]-tail[0]); L=18
        pts=[tip,(tip[0]-L*math.cos(ang-.55),tip[1]-L*math.sin(ang-.55)),(tip[0]-L*math.cos(ang+.55),tip[1]-L*math.sin(ang+.55))]
        d.polygon(pts,fill="#"+color)
    head(b,a)
    if both: head(a,b)


def label(d,pos,text,size=23,color=GRAY,bold=False,anchor="mm"):
    d.text(xy(pos),text,font=font(size,bold),fill="#"+color,anchor=anchor,align="center",spacing=7)


def save_image(im,name):
    im.save(IMG/name,dpi=(180,180),optimize=True)


def make_images():
    im,d=new_canvas("Figure 1. Coupling changes the total free-energy balance")
    pbox(d,.04,.45,.23,.22,"Endergonic step\nDelta G > 0",PALE_GOLD,GOLD)
    pbox(d,.39,.45,.23,.22,"ATP hydrolysis\nDelta G < 0",MINT,GREEN)
    pbox(d,.73,.45,.23,.22,"Coupled process\nDelta G(total) < 0",LIGHT,GREEN)
    arrow(d,(.28,.56),(.38,.56),GOLD); arrow(d,(.63,.56),(.72,.56),GREEN)
    label(d,(.50,.81),"Delta G(total) = sum of the coupled reaction Delta G values",24,GRAY)
    save_image(im,"bioenergetics.png")

    im,d=new_canvas("Figure 2. Glycolysis - invest, split, and collect")
    pbox(d,.03,.43,.18,.21,"Glucose",PALE_GOLD,GOLD)
    pbox(d,.28,.43,.20,.21,"6-carbon\nintermediates",MINT,GREEN)
    pbox(d,.55,.43,.16,.21,"2 x G3P",MINT,GREEN)
    pbox(d,.78,.43,.18,.21,"2 pyruvate\n+ 2 ATP net",PALE_GOLD,GOLD)
    for a,b in [((.22,.535),(.27,.535)),((.49,.535),(.54,.535)),((.72,.535),(.77,.535))]: arrow(d,a,b)
    label(d,(.17,.76),"2 ATP invested",22,RED,True); label(d,(.62,.76),"4 ATP + 2 NADH produced",22,GREEN,True)
    label(d,(.86,.87),"Low O2: pyruvate -> lactate; NADH -> NAD+",20,DARK)
    save_image(im,"glycolysis.png")

    im,d=new_canvas("Figure 3. Branching creates many enzyme-access points")
    cx,cy=xy((.36,.52)); d.ellipse((cx-34,cy-34,cx+34,cy+34),fill="#"+GOLD,outline="#"+DARK,width=4)
    for ang in [15,70,125,180,235,290,340]:
        pt=(.36+.25*math.cos(math.radians(ang)),.52+.25*math.sin(math.radians(ang)))
        arrow(d,(.36,.52),pt,GREEN,width=5); px,py=xy(pt); d.ellipse((px-22,py-22,px+22,py+22),fill="#"+MINT,outline="#"+GREEN,width=3)
    label(d,(.36,.86),"alpha(1->4) chains; alpha(1->6) branches",21,GRAY)
    pbox(d,.68,.43,.27,.18,"Many non-reducing ends\n= rapid mobilization",PALE_GOLD,GOLD,21)
    pbox(d,.68,.68,.27,.18,"Phosphorylase +\ndebranching enzyme",MINT,GREEN,21)
    save_image(im,"glycogen.png")

    im,d=new_canvas("Figure 4. Three bypasses make gluconeogenesis possible")
    ys=[.33,.53,.73]; left=["Pyruvate","F1,6BP","G6P"]; right=["PEP","F6P","Glucose"]
    enzymes=["Pyruvate carboxylase + PEPCK","Fructose-1,6-bisphosphatase","Glucose-6-phosphatase"]
    for y,l,r,e in zip(ys,left,right,enzymes):
        pbox(d,.03,y-.065,.18,.13,l,PALE_GOLD,GOLD,21); pbox(d,.79,y-.065,.17,.13,r,MINT,GREEN,21)
        arrow(d,(.22,y),(.78,y),GREEN); label(d,(.50,y-.045),e,20,DARK,True)
    label(d,(.50,.90),"Per glucose from 2 pyruvate: 4 ATP + 2 GTP + 2 NADH",22,RED,True)
    save_image(im,"gluconeogenesis.png")

    im,d=new_canvas("Figure 5. The hepatic fructose-2,6-bisphosphate switch")
    for y,vals,cols in [(.37,["Glucagon\nfasting","cAMP / PKA\nphosphorylation","F2,6BP\nfalls","Glycolysis down\nGNG up"],[PALE_GOLD,LIGHT,LIGHT,PALE_GOLD]),(.70,["Insulin\nfed","Phosphatase\ndephosphorylation","F2,6BP\nrises","Glycolysis up\nGNG down"],[MINT,LIGHT,LIGHT,MINT])]:
        xs=[.02,.27,.55,.78]; ws=[.18,.21,.17,.20]
        for x,w,v,fc in zip(xs,ws,vals,cols): pbox(d,x,y-.085,w,.17,v,fc,GREEN if y>.5 else GOLD,20)
        for a,b in [((.21,y),(.26,y)),((.49,y),(.54,y)),((.73,y),(.77,y))]: arrow(d,a,b,GREEN if y>.5 else GOLD)
    save_image(im,"regulation.png")

    im,d=new_canvas("Figure 6. Pentose phosphate pathway - two connected jobs")
    pbox(d,.04,.43,.18,.21,"Glucose-6-P",PALE_GOLD,GOLD)
    pbox(d,.31,.39,.25,.19,"Oxidative phase\n2 NADPH + CO2",MINT,GREEN)
    pbox(d,.66,.39,.28,.19,"Ribulose-5-P /\nRibose-5-P",LIGHT,GREEN)
    arrow(d,(.23,.535),(.30,.49)); arrow(d,(.57,.49),(.65,.49))
    pbox(d,.36,.69,.25,.17,"Non-oxidative\nreversible",PALE_GOLD,GOLD,21)
    pbox(d,.70,.69,.24,.17,"F6P + G3P\n(glycolysis)",LIGHT,GREEN,21)
    arrow(d,(.79,.59),(.58,.68),GOLD,True); arrow(d,(.62,.775),(.69,.775),GOLD,True)
    label(d,(.14,.80),"NADPH supports\nreduced glutathione",20,RED,True)
    save_image(im,"ppp.png")

    im,d=new_canvas("Figure 7. Glyoxylate bypass preserves carbon")
    pbox(d,.03,.45,.17,.19,"2 Acetyl-CoA",PALE_GOLD,GOLD)
    pbox(d,.28,.45,.17,.19,"Isocitrate",MINT,GREEN)
    pbox(d,.54,.34,.18,.17,"CAC route\n2 CO2 lost",LIGHT,RED,20)
    pbox(d,.54,.62,.18,.17,"Glyoxylate route\nno CO2 here",MINT,GREEN,20)
    pbox(d,.80,.62,.17,.17,"Succinate\n4-carbon output",PALE_GOLD,GOLD,20)
    arrow(d,(.21,.545),(.27,.545)); arrow(d,(.46,.52),(.53,.43),RED); arrow(d,(.46,.57),(.53,.69),GREEN); arrow(d,(.73,.705),(.79,.705),GREEN)
    label(d,(.64,.88),"Unique enzymes: isocitrate lyase + malate synthase",22,DARK,True)
    save_image(im,"cac_glyoxylate.png")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn("w:"+m))
        if node is None: node=OxmlElement("w:"+m); tcMar.append(node)
        node.set(qn("w:w"),str(v)); node.set(qn("w:type"),"dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); node = OxmlElement("w:tblHeader"); node.set(qn("w:val"),"true"); trPr.append(node)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    node.set(qn("w:val"), "true")
    trPr.append(node)


def set_font(run, size=11, bold=False, color=DARK, italic=False):
    run.font.name="Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)


def style_doc(doc, title):
    sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(DARK)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,before,after,color in [("Title",28,0,8,DARK),("Subtitle",13,0,14,GRAY),("Heading 1",16,18,10,GREEN),("Heading 2",13,14,7,GREEN),("Heading 3",12,10,5,DARK)]:
        st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.bold=(name!="Subtitle"); st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    header=sec.header.paragraphs[0]; header.text="SBIA022 | Carbohydrate Metabolism Challenge"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs: set_font(r,8.5,False,GRAY)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=footer.add_run(title); set_font(r,8.5,False,GRAY)


def add_cover(doc, teacher=False):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(38); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("METABOLIC MISSION"); set_font(r,30,True,GREEN)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("From ATP to the Glyoxylate Cycle"); set_font(r,17,True,DARK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Second-Year SBIA022 Quiz Pack" + (" - Lecturer Memo" if teacher else "")); set_font(r,12,False,GRAY)
    p.paragraph_format.space_after=Pt(25)
    t=doc.add_table(rows=4,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    labels=[("Institution","University of Limpopo"),("Question set","20 core questions + 20 challenge twins"),("Assessment","100 marks if both versions are completed"),("Progression","Remember -> Understand -> Apply -> Analyze -> Evaluate -> Create")]
    for i,(a,b) in enumerate(labels):
        t.rows[i].cells[0].width=Inches(1.65); t.rows[i].cells[1].width=Inches(4.65)
        set_cell_shading(t.rows[i].cells[0],MINT); set_cell_shading(t.rows[i].cells[1],LIGHT)
        for c in t.rows[i].cells: set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r=t.rows[i].cells[0].paragraphs[0].add_run(a); set_font(r,10,True,GREEN)
        r=t.rows[i].cells[1].paragraphs[0].add_run(b); set_font(r,10.5,False,DARK)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(22); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Think like a cell: conserve carbon, balance energy, and justify every pathway choice."); set_font(r,12,True,GOLD,True)
    if not teacher:
        for label in ["Student name: __________________________________________", "Student number: _______________________________________", "Date: ____________________   Tutorial group: _____________"]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); r=p.add_run(label); set_font(r,10.5,False,DARK)
    else:
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20)
        r=p.add_run("Memo use: "); set_font(r,10.5,True,RED)
        r=p.add_run("Accept scientifically equivalent wording. For higher-order questions, award marks for a correct causal chain, not for memorising this memo verbatim."); set_font(r,10.5,False,DARK)
    doc.add_page_break()


def add_instructions(doc, teacher=False):
    doc.add_heading("How to use this pack",1)
    items = [
        "Each numbered item contains a Core question (A) and a matched Challenge Twin (B). Together they make 40 questions.",
        "Questions within each topic move from lower-order recall toward application, analysis, evaluation, or creation.",
        "Show calculations and metabolic reasoning. A pathway name without an explanation may earn only partial credit.",
        "Figures are simplified learning aids; use the lecture terminology when answering.",
    ]
    if teacher: items.append("Suggested total: 100 marks - 50 for all A questions and 50 for all B questions. The twins can also be used as a second form, revision round, or make-up quiz.")
    for item in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.38); p.paragraph_format.first_line_indent=Inches(-.19); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25
        r=p.add_run(item); set_font(r,10.8,False,DARK)
    doc.add_heading("Bloom ladder and marks",2)
    p=doc.add_paragraph("Remember (1) | Understand (2) | Apply (3) | Analyze (3-4) | Evaluate (5) | Create (5)")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: set_font(r,10.5,True,GREEN)


def add_topic_header(doc, topic):
    doc.add_page_break()
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(topic["tag"]); set_font(r,9,True,GOLD)
    doc.add_heading(topic["title"],1)
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    prevent_row_split(t.rows[0])
    t.rows[0].cells[0].width=Inches(6.5); set_cell_shading(t.rows[0].cells[0],MINT); set_cell_margins(t.rows[0].cells[0],140,180,140,180)
    p=t.rows[0].cells[0].paragraphs[0]; r=p.add_run("MISSION BRIEF\n"); set_font(r,9,True,GREEN)
    r=p.add_run(topic["scenario"]); set_font(r,10.5,False,DARK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3); p.paragraph_format.keep_with_next=True
    r=p.add_run(); shape=r.add_picture(str(IMG/topic["image"]),width=Inches(6.15))
    shape._inline.docPr.set("title", topic["title"])
    shape._inline.docPr.set("descr", "Simplified study diagram for " + topic["title"])
    p=doc.add_paragraph(Path(topic["image"]).stem.replace("_"," ").title()+" learning schematic")
    p.style="Caption"; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
    for r in p.runs: set_font(r,8.5,False,GRAY,True)


def add_question(doc, q, teacher=False):
    p=doc.add_paragraph(); p.paragraph_format.keep_with_next=True; p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(f"Q{q['n']}  {q['level'].upper()}  |  {q['marks']} mark{'s' if q['marks'] != 1 else ''} each"); set_font(r,11,True,GREEN)
    for label,key,ak in [("A - Core","a","ans_a"),("B - Challenge Twin","b","ans_b")]:
        t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
        prevent_row_split(t.rows[0])
        cell=t.rows[0].cells[0]; cell.width=Inches(6.5); set_cell_shading(cell, LIGHT if label.startswith("A") else PALE_GOLD); set_cell_margins(cell,100,150,100,150)
        p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2)
        r=p.add_run(label+": "); set_font(r,10.5,True,GREEN if label.startswith("A") else GOLD)
        r=p.add_run(q[key]); set_font(r,10.5,False,DARK)
        if teacher:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.space_after=Pt(5)
            r=p.add_run("Suggested answer: "); set_font(r,9.5,True,RED)
            r=p.add_run(q[ak]); set_font(r,9.5,False,DARK)
        else:
            lines = 1 if q["marks"] == 1 else 2 if q["marks"] <= 3 else 3
            for _ in range(lines):
                p=doc.add_paragraph("________________________________________________________________________________")
                p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(1)
                for r in p.runs: set_font(r,8,False,"B7C0BD")


def add_source_note(doc):
    doc.add_page_break(); doc.add_heading("Source lectures used",1)
    sources=[
        "Bioenergetics",
        "Glycolysis",
        "Glycogen Metabolism",
        "Gluconeogenesis",
        "Regulation of Carbohydrate Metabolism",
        "Pentose Phosphate Pathway",
        "CAC Regulation and the Glyoxylate Cycle",
    ]
    p=doc.add_paragraph("This quiz was developed from the supplied SBIA022 lecture documents by Dr KW Poopedi and colleagues. The diagrams are simplified original study schematics prepared for this quiz.")
    for s in sources:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.38); p.paragraph_format.first_line_indent=Inches(-.19); r=p.add_run(s); set_font(r,10.5,False,DARK)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(15); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("END OF METABOLIC MISSION"); set_font(r,14,True,GREEN)


def build(path, teacher=False):
    doc=Document(); style_doc(doc, "Lecturer Memo" if teacher else "Student Quiz")
    add_cover(doc,teacher); add_instructions(doc,teacher)
    for topic in topics:
        add_topic_header(doc,topic)
        for q in topic["questions"]: add_question(doc,q,teacher)
    add_source_note(doc)
    props=doc.core_properties
    props.title="Metabolic Mission: SBIA022 Carbohydrate Metabolism Quiz"
    props.subject="Second-year University of Limpopo metabolism quiz"
    props.author="Prepared from supplied SBIA022 lecture materials"
    props.keywords="bioenergetics, glycolysis, glycogen, gluconeogenesis, PPP, CAC, glyoxylate"
    doc.save(path)


if __name__ == "__main__":
    make_images()
    build(OUT/"SBIA022_Metabolic_Mission_Student_Quiz.docx",False)
    build(OUT/"SBIA022_Metabolic_Mission_Lecturer_Memo.docx",True)
    print("Created quiz documents in",OUT)
